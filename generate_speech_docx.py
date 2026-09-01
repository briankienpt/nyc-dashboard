import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_presentation_doc(output_path):
    doc = docx.Document()

    # Page Margins (Normal 1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37) # Dark slate
    normal_style.paragraph_format.line_spacing = 1.2
    normal_style.paragraph_format.space_after = Pt(6)

    # Header Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("KỊCH BẢN THUYẾT TRÌNH BẢO VỆ ĐỒ ÁN TỐT NGHIỆP")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Navy Blue
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(14)
    sub_run = sub_p.add_run("Đề tài: Hệ thống BI Dashboard Phân tích Thị trường Bất động sản New York (BDS NYC)\nSlide: TỔNG QUAN THỊ TRƯỜNG")
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_section_header(text, icon="📌"):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run(f"{icon} {text}")
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        return h

    def add_speech_paragraph(text, italic=False, bold_prefix="", highlight=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.font.bold = True
            r_pre.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.italic = italic
        if highlight:
            r.font.color.rgb = RGBColor(0x04, 0x78, 0x57) # Green dark
        return p

    def add_callout(step_title, text):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        cell.width = Inches(6.5)
        
        # Light grey-blue background and left border
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F3F4F6"/>')
        borders = parse_xml(f'''
            <w:tcBorders {nsdecls("w")}>
                <w:top w:val="none"/>
                <w:left w:val="single" w:sz="24" w:space="0" w:color="3B82F6"/>
                <w:bottom w:val="none"/>
                <w:right w:val="none"/>
            </w:tcBorders>
        ''')
        tcPr.append(shd)
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.right_indent = Inches(0.1)
        
        r_title = p.add_run(f"👉 {step_title}\n")
        r_title.font.bold = True
        r_title.font.size = Pt(10.5)
        r_title.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        
        r_txt = p.add_run(f'"{text}"')
        r_txt.font.italic = True
        r_txt.font.size = Pt(11)
        r_txt.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 1. MỞ ĐẦU
    add_section_header("PHẦN 1. MỞ ĐẦU TRANG TỔNG QUAN", "🎙️")
    add_speech_paragraph(
        "Kính thưa quý Thầy Cô trong Hội đồng, bước vào phần cốt lõi của đề tài, em xin phép được trình bày trang đầu tiên của Dashboard: Trang Tổng quan Thị trường. "
        "Mục tiêu trọng tâm của trang này là giúp người dùng nắm bắt ngay lập tức quy mô toàn cảnh, mặt bằng định giá và cơ cấu nền tảng của thị trường bất động sản New York trước khi đi sâu vào các phân tích chi tiết."
    )
    add_speech_paragraph(
        "Nhìn lên các thẻ chỉ số KPI tổng hợp phía trên, tập dữ liệu ghi nhận hơn 47.000 giao dịch với tổng giá trị thị trường đạt hơn 47,1 tỷ USD. "
        "Mức giá trung bình toàn thành phố nằm ở mốc 865.000 USD, trong đó có tới gần 38,5% số giao dịch đạt giá trị từ 1 triệu USD trở lên — cho thấy New York là một thị trường có sức hút và tỷ trọng giá trị cao rất rõ rệt."
    )

    # 2. BIỂU ĐỒ 1
    add_section_header("BIỂU ĐỒ 1: SỐ LƯỢNG GIAO DỊCH THEO 5 QUẬN (THANH KHOẢN KHU VỰC)", "📊")
    add_callout("Lời thoại thuyết trình (Biểu đồ 1)",
        "Ở góc trên bên trái, biểu đồ đầu tiên thể hiện thanh khoản của thị trường thông qua số lượng giao dịch tại 5 quận. "
        "Điểm nổi bật nhất ở đây là Queens dẫn đầu toàn thành phố với hơn 14.400 giao dịch, theo sát là Manhattan với hơn 12.800 và Brooklyn với gần 12.000 giao dịch. Trong khi đó, Staten Island và Bronx có khối lượng giao dịch khiêm tốn hơn hẳn, chỉ từ 3.500 đến 4.200 giao dịch. "
        "Kết quả này phản ánh rằng hoạt động mua bán tại New York tập trung chủ yếu vào 'tam giác' Queens, Manhattan và Brooklyn — đây chính là 3 trung tâm thanh khoản cốt lõi của toàn thành phố. "
        "Tuy nhiên, nếu chỉ nhìn vào số lượng giao dịch thì chưa thể phản ánh giá trị thực tế của từng khu vực. Vì vậy, chúng ta cùng nhìn sang biểu đồ ngay bên cạnh để đánh giá về mặt bằng giá."
    )

    # 3. BIỂU ĐỒ 2
    add_section_header("BIỂU ĐỒ 2: GIÁ TRUNG BÌNH THEO QUẬN (MẶT BẰNG ĐỊNH GIÁ)", "📊")
    add_callout("Lời thoại thuyết trình (Biểu đồ 2)",
        "Tại biểu đồ thứ hai về Giá trung bình theo quận, em sử dụng chỉ số trung bình (Median) thay vì giá trung bình nhằm loại trừ hoàn toàn ảnh hưởng của các bất động sản siêu đắt đỏ. "
        "Kết quả hiển thị một bức tranh rất rõ nét: Manhattan và Brooklyn hoàn toàn áp đảo với mức giá trung bình vượt mốc 1 triệu USD (cụ thể Manhattan là 1,055 triệu USD và Brooklyn là 1,05 triệu USD). Con số này cao hơn gấp 1,4 đến 1,5 lần so với Queens (745.000 USD), Staten Island (720.000 USD) và Bronx (700.000 USD). "
        "Điều này chứng minh thị trường có sự phân tầng giá trị rất sâu sắc: Manhattan và Brooklyn đại diện cho phân khúc bất động sản cao cấp, giá trị lớn; trong khi Queens và Bronx lại đóng vai trò là thị trường nhà ở với mức giá dễ tiếp cận hơn cho đa số người dân. "
        "Để giải mã lý do tại sao có sự chênh lệch lớn này, biểu đồ tiếp theo sẽ bóc tách cơ cấu các loại hình bất động sản được giao dịch trên thị trường."
    )

    # 4. BIỂU ĐỒ 3
    add_section_header("BIỂU ĐỒ 3: CƠ CẤU LOẠI HÌNH BẤT ĐỘNG SẢN (DONUT CHART)", "📊")
    add_callout("Lời thoại thuyết trình (Biểu đồ 3)",
        "Biểu đồ tròn ở khu vực giữa slide thể hiện tỷ trọng giao dịch của các loại hình bất động sản phổ biến. "
        "Qua kết quả này, có thể thấy thị trường New York bị chi phối gần như tuyệt đối bởi hai nhóm chính: Căn hộ chung cư có thang máy (Elevator Apartments) chiếm 40,8% và Nhà ở gia đình (Family Dwellings) chiếm 40,1%. Tổng cộng hai loại hình này đã chiếm tới hơn 80% toàn bộ giao dịch của thị trường, trong khi các căn hộ Walkup chỉ chiếm 8,8% và các phân khúc khác chiếm tỷ lệ rất nhỏ. "
        "Cơ cấu này phản ánh đúng đặc trưng đô thị hóa của New York: mật độ dân cư cao thúc đẩy nhu cầu cực lớn về căn hộ chung cư cao tầng, song song với nhu cầu nhà ở gia đình truyền thống tại các quận ngoại vi. "
        "Vậy mức giá của từng loại hình này đang dao động như thế nào? Em xin mời Hội đồng nhìn sang biểu đồ Boxplot phân bố giá kế bên."
    )

    # 5. BIỂU ĐỒ 4
    add_section_header("BIỂU ĐỒ 4: PHÂN BỐ GIÁ THEO LOẠI HÌNH BẤT ĐỘNG SẢN (BOX PLOT TOP 6)", "📊")
    add_callout("Lời thoại thuyết trình (Biểu đồ 4)",
        "Biểu đồ phân bố giá Boxplot cho phép chúng ta quan sát cả mức giá trung bình lẫn biên độ dao động của top 6 loại hình bất động sản. "
        "Điểm đáng chú ý là nhóm Nhà ở phức hợp 2-10 căn (2-10 Unit Residential) có mức giá trung bình cao nhất, đạt khoảng 1,1 triệu USD, theo sau là Nhà ở gia đình (Family Dwellings) ở mức 900.000 USD. Trong khi đó, nhóm Căn hộ thang máy dù chiếm số lượng giao dịch lớn nhất nhưng mức giá trung bình rơi vào khoảng 760.000 USD với dải phân bố rất rộng. "
        "Ý nghĩa thực tiễn ở đây là các bất động sản tạo ra dòng tiền hoặc có diện tích sử dụng lớn luôn đòi hỏi vốn đầu tư ban đầu cao hơn, trong khi căn hộ chung cư cung cấp nhiều lựa chọn đa dạng từ phân khúc tầm trung đến cao cấp. "
        "Tiếp theo, để hiểu rõ hơn về đối tượng tham gia thị trường, chúng ta hãy xem xét cơ cấu phân khúc khách hàng ở phần bên dưới."
    )

    # 6. BIỂU ĐỒ 5 & 6
    add_section_header("BIỂU ĐỒ 5 & 6: CƠ CẤU PHÂN KHÚC KHÁCH HÀNG & MẶT BẰNG GIÁ THEO PHÂN KHÚC", "📊")
    add_callout("Lời thoại thuyết trình (Biểu đồ 5 & 6)",
        "Ở cụm biểu đồ phân khúc khách hàng, dữ liệu được phân chia dựa trên quy mô số căn trong tòa nhà nhằm phân loại mục đích mua: Mua ở thực (1 căn), Đầu tư nhỏ lẻ (2 đến 10 căn) và Nhà đầu tư tổ chức (trên 10 căn). "
        "Số liệu chỉ ra rằng nhóm Mua ở thực chiếm tỷ trọng áp đảo lên tới gần 95% tổng giao dịch toàn thị trường với mức giá trung bình khoảng 850.000 USD. Ngược lại, nhóm đầu tư từ 2 căn trở lên chỉ chiếm khoảng 5% nhưng mặt bằng giá trung bình lại tăng vọt lên từ 1,1 triệu đến hơn 2,5 triệu USD. "
        "Điều này khẳng định thị trường bất động sản New York về căn bản vẫn được nâng đỡ vững chắc bởi nhu cầu an cư thực tế của người dân, trong khi dòng vốn đầu tư tổ chức tập trung vào các tài sản quy mô lớn với giá trị giao dịch cao."
    )

    # 7. Ý NGHĨA TỔNG QUAN
    add_section_header("Ý NGHĨA TỔNG QUAN CỦA TRANG DASHBOARD (TỔNG KẾT BỨC TRANH THỊ TRƯỜNG)", "🎯")
    add_speech_paragraph(
        "Kính thưa Hội đồng, kết nối tất cả các biểu đồ trên lại với nhau, trang Tổng quan này kể cho chúng ta một câu chuyện rất hoàn chỉnh và nhất quán về thị trường bất động sản New York:", bold_prefix="📌 "
    )
    add_speech_paragraph(
        "Thứ nhất, thị trường có quy mô cực lớn (hơn 47 tỷ USD) và được duy trì thanh khoản mạnh mẽ bởi nhu cầu ở thực của người dân (chiếm 95% giao dịch, tập trung vào Queens, Manhattan và Brooklyn). "
        "Thứ hai, thị trường có sự phân hóa không gian vô cùng sâu sắc: Manhattan và Brooklyn dẫn dắt mặt bằng giá cao cấp (trên 1 triệu USD), trong khi Queens và Bronx đóng vai trò phân khúc giá đại chúng. "
        "Và thứ ba, cấu trúc sản phẩm định hình rõ nét với hơn 80% giao dịch nằm ở căn hộ chung cư cao tầng và nhà ở gia đình."
    )
    add_speech_paragraph(
        "Như vậy, trang Tổng quan đã hoàn thành xuất sắc vai trò cung cấp cho nhà đầu tư và người dùng một bức tranh toàn cảnh chính xác, đa chiều và đáng tin cậy về thị trường NYC trước khi đi vào các phân tích chuyên sâu.", bold_prefix="✨ "
    )

    # 8. CÂU CHUYỂN
    add_section_header("CÂU CHUYỂN TIẾP TỰ NHIÊN SANG SLIDE TIẾP THEO", "⏭️")
    add_callout("Câu chuyển kết thúc slide",
        "Và để hiểu rõ hơn từng khu vực cụ thể đang có mức giá biến động ra sao và đâu là những 'điểm nóng' bất động sản thực tế trên bản đồ không gian, em xin kính mời quý Thầy Cô cùng chuyển sang slide tiếp theo: Phân tích Chi tiết Khu vực và Bản đồ Nhiệt Hotspot Map."
    )

    doc.save(output_path)
    print(f"File saved successfully at {output_path}")

create_presentation_doc("c:/Users/phong/OneDrive/Desktop/NYC_Dashboard/Loi_Thuyet_Trinh_Slide_Tong_Quan.docx")
